def create_report(job):

    global os
    global pd
    global interpolate
    global plt
    global np
    global rcParams
    global der
    global Document
    global MailMerge
    global mn
    global cur_dir
    
    import pandas as pd
    import os
    import numpy as np
    import matplotlib.pyplot as plt
    from matplotlib import rcParams
    from scipy import interpolate
    from scipy.misc import derivative as der
    from docx import Document
    from mailmerge import MailMerge
    from menu import menu as mn
    
    cur_dir = os.getcwd() + '\\'

    si = get_str('SI#')
    tp = choose_op(('Dedication', 'Post Rewind','Non-Specific'))
    volt = get_float('required test voltage')
    per_volt = choose_st()
    filepath = find_file(per_volt, job, 'st')
    filepath_lr = find_file(per_volt, job, 'lr')
    
    df = pd.read_csv(filepath)
    df_lr = pd.read_csv(filepath_lr)
    
    df = clean_data(df, df_lr, volt)
    plot_data(df)
    
    
    make_word_file(df, job, si, per_volt, tp)
    #export_excel(df_f)

def export_excel(df_f):

    export_to = r'C:\Users\kenneym\Desktop\speed_torq_files\table.xlsx'
    with pd.ExcelWriter(export_to) as writer:
        df_f.to_excel(writer)
    os.startfile(export_to)
    return

def clean_data(df, df_lr, volt):

    def round_data(data_frame):

        keys = data_frame.columns
        for i in range(data_frame.shape[0]):
            for j in keys:
                n = data_frame.at[i, j]
                tj = j.lower()
                if 'speed' in tj:
                    data_frame.at[i, j] = round(abs(n))
                elif 'torque' in tj:
                    data_frame.at[i, j] = round(abs(n),1)
        return data_frame

    #remove NaN entries
    df = df.dropna()
    df_lr = df_lr.dropna()

    #convert data types to float or integers, remove negative entries and
    #round data points
    df_lr = round_data(df_lr)
    df = round_data(df)
    
    #finding row with locked rotor data
    df_lr = df_lr[df_lr['Speed (RPM)'] == 0]
    df_lr.drop_duplicates()
    lrI = df_lr[['Torque (lbf-ft)']].idxmax()
    df_lr = pd.DataFrame(df_lr.iloc[lrI].to_dict()) 
    df_lr.reset_index(drop=True, inplace=True)

    #calculate locked rotor
    df_c = df.drop_duplicates(subset='Speed (RPM)', keep='last')
    ys = df_c['Torque (lbf-ft)'].tolist()
    xs = df_c['Speed (RPM)'].tolist()
    f1 = interpolate.interp1d(xs, ys, kind='cubic', fill_value='extrapolate')
    ds = 100 * der(f1, np.array(xs))
    ds = list(ds)
    ax1 = plt.plot(xs, ys), plt.plot(xs, ds)
    #plt.show()
    df = df[df['Speed (RPM)'] > xs[ds.index(max(ds))]]

    #appending locked rotor data
    df = df[df['Current Avg (RMS)'] > 1]
    df = df.append(df_lr)
    df.reset_index(inplace=True, drop=True)

    #correct data
    def conv_cur(vt, vr, ct):
        vt = abs(vt)
        ct = abs(ct)
        return round(ct * vt / vr, 1)

    df['Current 1 (RMS)'] = conv_cur(df['Volt 1 (RMS)'], volt, df['Current 1 (RMS)'])
    df['Current 2 (RMS)'] = conv_cur(df['Volt 2 (RMS)'], volt, df['Current 2 (RMS)'])
    df['Current 3 (RMS)'] = conv_cur(df['Volt 3 (RMS)'], volt, df['Current 3 (RMS)'])
    df['Current Avg (RMS)'] = conv_cur(df['Volt Avg (RMS)'], volt, df['Current Avg (RMS)'])
    df['Volt 1 (RMS)'] = volt
    df['Volt 2 (RMS)'] = volt
    df['Volt 3 (RMS)'] = volt
    df['Volt Avg (RMS)'] = volt

    #removing duplicates
    df = df.drop_duplicates(subset='Current Avg (RMS)', keep='last')
    df = df.drop_duplicates(subset='Torque (lbf-ft)', keep='last')
    df = df.drop_duplicates(subset='Speed (RPM)', keep='last')

    df.reset_index(drop=True, inplace=True)

    return df
    
def plot_data(df):

        
    xs = list(df['Speed (RPM)'])
    ys = list(df['Torque (lbf-ft)'])
    cs = list(df['Current Avg (RMS)'])
    df.to_html(r'C:\Users\kenneym\Desktop\test.html')

    xss = np.linspace(min(xs), max(xs), 500)
    css = np.linspace(min(cs), max(cs), 500)
    yss = np.linspace(min(ys), max(ys), 500)
    f1 = interpolate.interp1d(xs, ys, kind = 'linear', fill_value='extrapolate')
    f2 = interpolate.interp1d(xs, cs, kind = 'cubic', fill_value='extrapolate')
    
    fig = plt.figure()
    ax1 = plt.subplot(1,1,1)
    ax1.plot(ys, xs, color='blue', label='100% Speed vs. Torque')

    ax2 = ax1.twinx()
    ax2.plot(ys, cs, linestyle= '--',
                      color='red', label='100% Current vs. Torque')
    ax1.set_xlabel('Torque (lb-ft)', weight='heavy')
    ax1.set_ylabel('Speed (RPM)', weight='heavy')
    ax2.set_ylabel('Average Current (amps)', weight='heavy')
    rcParams['font.family'] = 'Calibri'
    ax1.minorticks_on()
    ax2.minorticks_on()
    ax1.grid(True, which='major', linewidth=1)
    ax1.grid(True, which='minor', linewidth=0.5)
    fig.legend(loc='center')
    plt.savefig(cur_dir + 'st100.jpg')

    #plt.show()
    
def find_file(per_volt, job, t):

    def find_folder(dir_0, parm):
        dir_1 = dir_0
        folder_list = os.listdir(dir_1)
        flag = False
        for i in folder_list:
            if flag == False:
                if parm.lower() in i.lower():
                    flag = True
                    folder = i
                    dir_2 = dir_1 + '\\' + folder
        return dir_2

    directory = r'T:\QA\QADEPT\Job Folders\\' 
    
    d = find_folder(directory, job.customer_name)
    d = find_folder(d, job.sec_job_number)
    d = find_folder(d, 'load test')
    d = find_folder(d, 'raw')
    d = find_folder(d, str(per_volt) + t)

    return d

def choose_op(chs):

    chI = mn(chs)
    ch = chs[chI]
    if ch == 'Non-Specific':
        ch = ''
    return ch

def choose_st():

    msg = 'Enter the percentage of voltage for testing: '
    flag = False
    while flag == False:
        n = input(msg.rjust(50))
        try:
            n = int(n)
        except:
            print('n\Enter only numbers. ')
        else:
            flag = True
    return str(n) + '%'

def get_float(name):
    msg = 'Enter the ' + name + ': '
    flag = False
    while flag == False:
        n = input(msg.rjust(50))
        try:
            n = float(n)
        except:
            print('\nEnter only numbers. ')
        else:
            flag = True
    return n

def get_str(name):

    msg = 'Enter the ' + name + ': '
    msg_f = msg.rjust(50)
    a = input(msg_f)
    return a

def make_word_file(df, job, si, st, tp):

    def conv_cur(vt, vr, ct):
        return ct * vt / vr

    in_file = cur_dir + 'st_word_temp.docx'
    out_file = cur_dir + 'st_report.docx'
    pic_file = cur_dir + 'st100.jpg'

    varis = {'tp': tp,
             'sjn': job.sec_job_number,
             'cust': job.customer_name,
             'po': str(job.purchase_order),
             'hp': str(job.ratings.get('hp')),
             'idn': job.unique_id,
             'si': si,
             'st': str(st)}

    varis_rows = []
    for i in range(df.shape[0]):
        t_dict = {}
        t_dict['t'] = str(df.at[i, 'Torque (lbf-ft)'])
        t_dict['rpm'] = str(df.at[i, 'Speed (RPM)'])
        t_dict['v1'] = str(df.at[i, 'Volt 1 (RMS)'])
        t_dict['v2'] = str(df.at[i, 'Volt 2 (RMS)'])
        t_dict['v3'] = str(df.at[i, 'Volt 3 (RMS)'])
        t_dict['va'] = str(df.at[i, 'Volt Avg (RMS)'])
        t_dict['c1'] = str(df.at[i, 'Current 1 (RMS)'])
        t_dict['c2'] = str(df.at[i, 'Current 2 (RMS)'])
        t_dict['c3'] = str(df.at[i, 'Current 3 (RMS)'])
        t_dict['ca'] = str(df.at[i, 'Current Avg (RMS)'])
        varis_rows.append(t_dict)
             
    doc = MailMerge(in_file)
    doc.merge(**varis)
    doc.merge_rows('t', varis_rows)
    doc.write(out_file)

    doc = Document(out_file)
    for par in doc.paragraphs:
        if '<<image>>' in par.text:
            par.text = par.text.strip().replace('<<image>>', '')
            run = par.add_run()
            run.add_picture(pic_file)
    
    doc.save(out_file)
    os.startfile(out_file)
    
